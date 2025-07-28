from docx import Document
from docx.shared import Inches

def generate_consultation_doc(template_path, output_path, chantier, descriptif, tableau, visuels):
    doc = Document(template_path)

    # Remplir champ chantier en page 1
    for p in doc.paragraphs:
        if "REFERENCE CHANTIER" in p.text:
            p.text = f"REFERENCE CHANTIER : {chantier}"
            break

    # PAGE 2 – Quantitatif
    doc.add_page_break()
    doc.add_heading("Tableau Quantitatif", level=1)
    if tableau:
        headers = tableau[0].keys()
        table = doc.add_table(rows=1, cols=len(headers))
        for i, h in enumerate(headers):
            table.cell(0, i).text = str(h)
        for row in tableau:
            cells = table.add_row().cells
            for i, h in enumerate(headers):
                cells[i].text = str(row[h])
    else:
        doc.add_paragraph("Aucune donnée trouvée dans le DPGF.")

    # PAGE 3 – Visuels
    doc.add_page_break()
    doc.add_heading("🔍 Visuels des châssis extraits du DCE", level=1)
    if visuels:
        for img_path in visuels[:5]:  # On limite à 5 images pour éviter surcharge
            doc.add_picture(img_path, width=Inches(5.5))
    else:
        doc.add_paragraph("Aucun visuel trouvé dans les pièces jointes.")

    doc.save(output_path)# Génération Word
