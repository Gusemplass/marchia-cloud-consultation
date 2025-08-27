from fastapi import FastAPI, Response, Query, UploadFile, File, Form, HTTPException
from pydantic import BaseModel
from typing import List, Optional, Tuple, Dict
from io import BytesIO
import zipfile, csv, json, re, os

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.enum.text import WD_BREAK
from docx.shared import Pt

__VERSION__ = "2025-08-27-17"
TEMPLATE_PATH = "templates/fiche_demo_MARCHIA_full.docx"

app = FastAPI()

# ---------- Modèles ----------
class LigneQuantitative(BaseModel):
    rep: str
    dim: str
    typo: str
    perf: str
    qte: int
    pose: str
    commentaire: Optional[str] = ""

class FicheRequest(BaseModel):
    projet: str
    moa: str
    lot: str
    descriptif: str
    lignes: Optional[List[LigneQuantitative]] = None

# ---------- Helpers DOCX ----------
def find_paragraph(doc: Document, needles: List[str]) -> Optional[Paragraph]:
    for p in doc.paragraphs:
        for n in needles:
            if n in p.text:
                return p
    return None

def block_items(doc: Document):
    """Yield Paragraph/Table dans l'ordre d'apparition (de haut en bas)."""
    body = doc._element.body
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)

def remove_paragraph(p: Paragraph):
    p._element.getparent().remove(p._element)

def move_table_after_paragraph(table: Table, paragraph: Paragraph):
    tbl = table._tbl
    parent = tbl.getparent()
    parent.remove(tbl)
    paragraph._p.addnext(tbl)

def clear_table_body_keep_header(table: Table):
    """Supprime toutes les lignes sauf la première (entête)."""
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)

def zero_cell_spacing(table: Table):
    """Espace avant/après = 0 dans toutes les cellules (évite les blancs parasites)."""
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                pf = p.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)

def _has_page_break(par: Paragraph) -> bool:
    return any(r._r.xpath('.//w:br[@w:type="page"]') for r in par.runs)

def _has_section_break(par: Paragraph) -> bool:
    return bool(par._p.xpath('.//w:sectPr'))

def _norm(txt: str) -> str:
    return txt.replace("\u00A0", " ").replace("\t", " ").strip()

def cleanup_after_marker(marker_par: Paragraph, doc: Document) -> Optional[Table]:
    junk_prefixes = {
        "📌Rép.", "📐Dim.", "🧩Typo.", "🎯", "🏷", "🔧", "🧾",
        "Rép.", "Dim.", "Typo.", "Perf.", "Qté", "Pose", "Commentaire"
    }
    items = list(block_items(doc))
    try:
        idx = next(i for i, it in enumerate(items) if isinstance(it, Paragraph) and it._p is marker_par._p)
    except StopIteration:
        return None
    j = idx + 1
    first_table = None
    while j < len(items):
        it = items[j]
        if isinstance(it, Table):
            first_table = it
            break
        if isinstance(it, Paragraph):
            txt = _norm(it.text)
            if (txt == "" or any(txt.startswith(p) for p in junk_prefixes) or _has_page_break(it) or _has_section_break(it)):
                remove_paragraph(it)
                items.pop(j)
                continue
            else:
                break
        j += 1
    return first_table

def build_doc(req: FicheRequest) -> bytes:
    doc = Document(TEMPLATE_PATH)
    try:
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
    except Exception:
        pass
    for p in doc.paragraphs:
        if "{{projet}}" in p.text:
            p.text = p.text.replace("{{projet}}", req.projet)
        if "{{moa}}" in p.text:
            p.text = p.text.replace("{{moa}}", req.moa)
        if "{{lot}}" in p.text:
            p.text = p.text.replace("{{lot}}", req.lot)
    p_desc = find_paragraph(doc, ["[[DESCRIPTIF_CCTP]]", "{{DESCRIPTIF_CCTP}}"])
    if p_desc:
        p_desc.text = req.descriptif
    p_tbl = find_paragraph(doc, ["[[TABLEAU_QUANTITATIF]]", "{{TABLEAU_QUANTITATIF}}"])
    if p_tbl and req.lignes:
        p_tbl.text = ""
        run = p_tbl.add_run()
        run.add_break(WD_BREAK.PAGE)
        try:
            next_idx = doc.paragraphs.index(p_tbl) + 1
            if next_idx < len(doc.paragraphs):
                doc.paragraphs[next_idx].paragraph_format.page_break_before = False
        except Exception:
            pass
        existing_table = cleanup_after_marker(p_tbl, doc)
        headers = ["Rép.", "Dim.", "Typo.", "Perf. (Uw / Rw+Ctr)", "Qté", "Pose", "Commentaire"]
        if existing_table is None:
            dest_table = doc.add_table(rows=1, cols=len(headers))
            hdr = dest_table.rows[0].cells
            for i, h in enumerate(headers):
                hdr[i].text = h
            p_tbl._p.addnext(dest_table._tbl)
        else:
            dest_table = existing_table
            move_table_after_paragraph(dest_table, p_tbl)
            clear_table_body_keep_header(dest_table)
            hdr = dest_table.rows[0].cells
            for i, h in enumerate(headers[:len(hdr)]):
                hdr[i].text = h
        for L in (req.lignes or []):
            row = dest_table.add_row().cells
            row[0].text = L.rep
            row[1].text = L.dim
            row[2].text = L.typo
            row[3].text = L.perf
            row[4].text = str(int(L.qte))
            row[5].text = L.pose
            row[6].text = (L.commentaire or "").strip()
        try:
            dest_table.style = "Table Grid"
        except Exception:
            pass
        zero_cell_spacing(dest_table)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ---------- Helpers DCE ----------
KEYWORDS_QUANT = re.compile(r"(quant|dpgf|bpu|bordereau|dqe|estimatif)", re.I)

SYNONYMS: Dict[str, List[str]] = {
    "rep": ["rep", "repere", "repère", "ref", "réf", "reference", "référence", "code", "rep."],
    "dim": ["dim", "dimension", "dimensions", "l x h", "lxh", "l*h", "format"],
    "typo": ["typo", "typologie", "type", "designation", "désignation", "article", "libelle", "libellé", "description"],
    "perf": ["perf", "performance", "performances", "uw", "rw", "ctr", "ei", "a2p", "acoustique"],
    "qte": ["qte", "qté", "quantite", "quantité", "qty", "q.", "q", "quant."],
    "pose": ["pose", "mise en oeuvre", "mise en œuvre", "fourniture et pose", "f&p"],
    "commentaire": ["commentaire", "observations", "remarques", "note"],
}

def _norm_key(s: str) -> str:
    s = (s or "").strip().lower()
    s = s.replace("œ", "oe").replace("é", "e").replace("è", "e").replace("ê", "e").replace("à", "a").replace("î", "i")
    s = re.sub(r"[^a-z0-9\. ]", " ", s)
    s = re.sub(r"\s+", " ", s)
    return s

def _map_headers(headers: List[str]) -> Dict[str, Optional[int]]:
    idx = { _norm_key(h): i for i, h in enumerate(headers) }
    mapping = {k: None for k in SYNONYMS.keys()}
    for target, syns in SYNONYMS.items():
        for raw_h, i in idx.items():
            if raw_h in syns or any(raw_h.startswith(s) for s in syns):
                mapping[target] = i
                break
    return mapping

def _extract_perf_from_text(txt: str) -> str:
    if not txt:
        return ""
    hits = []
    m = re.findall(r"\bEI\s*[\d]{1,2}\b", txt, flags=re.I)
    if m: hits.append(", ".join(sorted(set(m))))
    m = re.findall(r"Rw\+?Ctr\s*[=≥]?\s*[\d]{1,2}\s*dB", txt, flags=re.I)
    if m: hits.append(", ".join(sorted(set(m))))
    m = re.findall(r"Uw\s*[=≤]?\s*[\d][\.,]\d+\s*W", txt, flags=re.I)
    if m: hits.append(", ".join(sorted(set(m))))
    return " / ".join(hits)

def _extract_dim_from_text(txt: str) -> str:
    if not txt:
        return ""
    m = re.search(r"\b(\d{3,4})\s*[xX]\s*(\d{3,4})\b", txt)
    if m:
        return f"{m.group(1)}x{m.group(2)}"
    return ""

def _guess_meta_from_names(zipname: str, names: List[str]) -> Tuple[str, str]:
    base = os.path.splitext(os.path.basename(zipname))[0]
    projet = re.sub(r"[_\-]+", " ", base).strip().title() or "Projet DCE"
    lot = ""
    for n in names:
        m = re.search(r"\blot\s*(?:n[oº°]?\s*\d+\s*)?([a-z0-9 \-_]+)", n, flags=re.I)
        if m:
            lot = m.group(0)
            break
    lot = (lot.strip().title() or "Lot Non Précisé")
    return projet, lot

def _read_csv_quant(raw: str) -> List[LigneQuantitative]:
    try:
        dialect = csv.Sniffer().sniff(raw[:4096], delimiters=",;")
        delim = dialect.delimiter
    except Exception:
        delim = ","
    reader = csv.reader(raw.splitlines(), delimiter=delim)
    rows = list(reader)
    if not rows:
        return []
    headers = rows[0]
    mapping = _map_headers(headers)
    has_design = mapping["typo"] is not None or any(_norm_key(h).startswith(("designation","description")) for h in headers)
    has_qty = mapping["qte"] is not None or any(_norm_key(h).startswith("quantite") or _norm_key(h) in ("qte","q") for h in headers)
    lignes: List[LigneQuantitative] = []
    for i, r in enumerate(rows[1:], start=1):
        cells = r + [""] * max(0, len(headers) - len(r))
        def val(idx_opt, default=""):
            return (cells[idx_opt] if idx_opt is not None and idx_opt < len(cells) else default).strip()
        if not any(c.strip() for c in cells):
            continue
        if mapping["typo"] is None and has_design:
            try:
                j = next(j for j,h in enumerate(headers) if _norm_key(h).startswith(("designation","description")))
                mapping["typo"] = j
            except StopIteration:
                pass
        if mapping["qte"] is None and has_qty:
            try:
                j = next(j for j,h in enumerate(headers) if _norm_key(h).startswith("quantite") or _norm_key(h) in ("qte","q"))
                mapping["qte"] = j
            except StopIteration:
                pass
        typo = val(mapping["typo"]) if mapping["typo"] is not None else ""
        qraw = val(mapping["qte"]) if mapping["qte"] is not None else ""
        if not (typo and (qraw or qraw=="0")):
            continue
        try:
            qte = int(float(str(qraw).replace(",", ".").strip()))
        except Exception:
            qte = 0
        rep = val(mapping["rep"]) if mapping["rep"] is not None else f"L{i}"
        dim = val(mapping["dim"]) if mapping["dim"] is not None else _extract_dim_from_text(typo)
        perf = val(mapping["perf"]) if mapping["perf"] is not None else _extract_perf_from_text(typo)
        pose = val(mapping["pose"]) if mapping["pose"] is not None else ""
        com  = val(mapping["commentaire"]) if mapping["commentaire"] is not None else ""
        lignes.append(LigneQuantitative(rep=rep, dim=dim, typo=typo, perf=perf, qte=qte, pose=pose, commentaire=com))
    return lignes

def _try_read_xlsx_quant(data: bytes) -> List[LigneQuantitative]:
    try:
        import openpyxl  # type: ignore
    except Exception:
        return []
    wb = openpyxl.load_workbook(BytesIO(data), data_only=True)
    lignes: List[LigneQuantitative] = []
    for ws in wb.worksheets:
        header_row = None
        for r in ws.iter_rows(min_row=1, max_row=20, values_only=True):
            vals = [str(c).strip() for c in (r or []) if c not in (None, "")]
            if len(vals) >= 2:
                header_row = r
                break
        if not header_row:
            continue
        headers = [str(c or "").strip() for c in header_row]
        mapping = _map_headers(headers)
        has_design = mapping["typo"] is not None or any(_norm_key(h).startswith(("designation","description")) for h in headers)
        has_qty = mapping["qte"] is not None or any(_norm_key(h).startswith("quantite") or _norm_key(h) in ("qte","q") for h in headers)
        if not (has_design and has_qty):
            continue
        start_row = 2
        for i in range(start_row, ws.max_row + 1):
            row = [ws.cell(row=i, column=j+1).value for j in range(len(headers))]
            cells = [str(c or "").strip() for c in row]
            if all(c == "" for c in cells):
                continue
            def val(idx_opt, default=""):
                return (cells[idx_opt] if idx_opt is not None and idx_opt < len(cells) else default).strip()
            if mapping["typo"] is None and has_design:
                try:
                    j = next(j for j,h in enumerate(headers) if _norm_key(h).startswith(("designation","description")))
                    mapping["typo"] = j
                except StopIteration:
                    pass
            if mapping["qte"] is None and has_qty:
                try:
                    j = next(j for j,h in enumerate(headers) if _norm_key(h).startsWith("quantite") or _norm_key(h) in ("qte","q"))
                except Exception:
                    pass
            typo = val(mapping["typo"]) if mapping["typo"] is not None else ""
            qraw = val(mapping["qte"]) if mapping["qte"] is not None else ""
            if not (typo and (qraw or qraw=="0")):
                continue
            try:
                qte = int(float(str(qraw).replace(",", ".").strip()))
            except Exception:
                qte = 0
            rep = val(mapping["rep"]) if mapping["rep"] is not None else f"L{i-1}"
            dim = val(mapping["dim"]) if mapping["dim"] is not None else _extract_dim_from_text(typo)
            perf = val(mapping["perf"]) if mapping["perf"] is not None else _extract_perf_from_text(typo)
            pose = val(mapping["pose"]) if mapping["pose"] is not None else ""
            com  = val(mapping["commentaire"]) if mapping["commentaire"] is not None else ""
            lignes.append(LigneQuantitative(rep=rep, dim=dim, typo=typo, perf=perf, qte=qte, pose=pose, commentaire=com))
        if lignes:
            break
    return lignes

def _find_quant_file(names: List[str]) -> Optional[str]:
    candidates = []
    for n in names:
        low = n.lower()
        if low.endswith((".csv", ".xlsx")) and KEYWORDS_QUANT.search(low):
            weight = 100
        elif low.endswith((".csv", ".xlsx")):
            weight = 10
        else:
            continue
        size_bias = -len(n)
        candidates.append((weight, size_bias, n))
    if not candidates:
        for n in names:
            if n.lower().endswith((".csv", ".xlsx")):
                candidates.append((1, -len(n), n))
    if not candidates:
        return None
    candidates.sort(reverse=True)
    return candidates[0][2]

# ---------- Routes ----------
@app.get("/")
def root():
    return {"message": "Marchia Cloud Consultation en ligne", "version": __VERSION__}

@app.get("/health")
def health():
    return {"ok": True, "version": __VERSION__}

@app.post("/genere-fiche")
def genere_fiche(req: FicheRequest, format: str = Query("docx", pattern="^(docx|json)$")):
    if format == "json":
        return {"status": "ok", "message": "Fiche reçue correctement", "data": req.model_dump()}
    content = build_doc(req)
    filename = f'fiche_{req.projet.replace(" ", "_")}.docx'
    return Response(
        content=content,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )

@app.post("/genere-fiche-zip")
async def genere_fiche_zip(
    file: UploadFile = File(..., description="ZIP avec quantitatif.csv (+ meta.json optionnel)"),
    projet: Optional[str] = Form(None),
    moa: Optional[str] = Form(None),
    lot: Optional[str] = Form(None),
    descriptif: Optional[str] = Form(None),
):
    data = await file.read()
    try:
        zf = zipfile.ZipFile(BytesIO(data))
    except zipfile.BadZipFile:
        raise HTTPException(status_code=400, detail="Fichier non valide: ZIP attendu.")
    names = zf.namelist()
    qcsv_name = next((n for n in names if n.lower().endswith("quantitatif.csv")), None)
    if not qcsv_name:
        raise HTTPException(status_code=400, detail="quantitatif.csv introuvable dans le ZIP.")
    meta_name = next((n for n in names if n.lower().endswith("meta.json")), None)
    meta = {}
    if meta_name:
        try:
            meta = json.loads(zf.read(meta_name).decode("utf-8-sig"))
        except Exception:
            meta = {}
    _projet = (projet or meta.get("projet") or "").strip()
    _moa = (moa or meta.get("moa") or "").strip()
    _lot = (lot or meta.get("lot") or "").strip()
    _desc = (descriptif or meta.get("descriptif") or "").strip()
    if not (_projet and _moa and _lot):
        raise HTTPException(status_code=400, detail="Champs requis manquants (projet, moa, lot).")
    try:
        raw = zf.read(qcsv_name).decode("utf-8-sig")
    except UnicodeDecodeError:
        raw = zf.read(qcsv_name).decode("latin-1")
    lignes = _read_csv_quant(raw)
    if not lignes:
        raise HTTPException(status_code=400, detail="Aucune ligne exploitable trouvée dans quantitatif.csv.")
    req = FicheRequest(projet=_projet, moa=_moa, lot=_lot, descriptif=_desc, lignes=lignes)
    content = build_doc(req)
    filename = f'fiche_{_projet.replace(" ", "_")}.docx'
    return Response(
        content=content,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )

@app.post("/genere-fiche-dce")
async def genere_fiche_dce(
    file: UploadFile = File(..., description="ZIP DCE brut (PDF/Excel/CSV)"),
    projet: Optional[str] = Form(None),
    moa: Optional[str] = Form(None),
    lot: Optional[str] = Form(None),
    descriptif: Optional[str] = Form(None),
):
    data = await file.read()
    try:
        zf = zipfile.ZipFile(BytesIO(data))
    except zipfile.BadZipFile:
        raise HTTPException(status_code=400, detail="Fichier non valide: ZIP attendu.")
    names = zf.namelist()
    quant_name = _find_quant_file(names)
    if not quant_name:
        raise HTTPException(status_code=400, detail="Aucun fichier quantitatif (.csv/.xlsx) détecté (cherché: quant, dpgf, bpu, dqe, bordereau, estimatif).")
    lignes: List[LigneQuantitative] = []
    if quant_name.lower().endswith(".csv"):
        try:
            raw = zf.read(quant_name).decode("utf-8-sig")
        except UnicodeDecodeError:
            raw = zf.read(quant_name).decode("latin-1")
        lignes = _read_csv_quant(raw)
    elif quant_name.lower().endswith(".xlsx"):
        lignes = _try_read_xlsx_quant(zf.read(quant_name))
    if not lignes:
        raise HTTPException(status_code=400, detail=f"Quantitatif '{os.path.basename(quant_name)}' non exploitable (désignation/quantité manquantes ?).")
    meta_name = next((n for n in names if n.lower().endswith("meta.json")), None)
    meta = {}
    if meta_name:
        try:
            meta = json.loads(zf.read(meta_name).decode("utf-8-sig"))
        except Exception:
            meta = {}
    _projet = (projet or meta.get("projet") or "").strip()
    _moa = (moa or meta.get("moa") or "").strip()
    _lot = (lot or meta.get("lot") or "").strip()
    _desc = (descriptif or meta.get("descriptif") or "").strip()
    if not _projet or not _lot:
        g_proj, g_lot = _guess_meta_from_names(file.filename or "DCE.zip", names)
        _projet = _projet or g_proj
        _lot = _lot or g_lot
    if not _moa:
        _moa = "MOA non précisée"
    req = FicheRequest(projet=_projet, moa=_moa, lot=_lot, descriptif=_desc, lignes=lignes)
    content = build_doc(req)
    filename = f'fiche_{_projet.replace(" ", "_")}.docx'
    return Response(
        content=content,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )
